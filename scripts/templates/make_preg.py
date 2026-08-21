# 임신기 근로시간 단축근무 신청서 — 필드 심기 (매니저용 / 원장용)
#
# 두 문서는 결재표만 다르고 나머지는 동일하다.
#   매니저용 결재표: 작성자 | 원장 | 본부
#   원장용   결재표: 작성자 | 본부
#
# 심는 것
#   자동    {직원명} {생년월일} {지점} {직급} {작성일}
#   관리자  {체크_12주이내} {체크_32주이후} {체크_기타}
#           {출산예정일} {단축시작일} {단축종료일} {근무개시시각} {근무종료시각}
#   서명    《근로자서명》 《원장서명》 《본부서명》
import re, sys, shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = {
    "manager": r"C:/Users/N-88/Downloads/직영점_임신기 근로시간 단축근무 신청서 - 매니저용.docx",
    "director": r"C:/Users/N-88/Downloads/직영점_임신기 근로시간 단축근무 신청서 - 원장용.docx",
}
OUT = {
    "manager": "임신기_단축근무_매니저용_템플릿.docx",
    "director": "임신기_단축근무_원장용_템플릿.docx",
}


def cell_text(c):
    return "".join(p.text for p in c.paragraphs)


def set_cell(cell, text, align=None):
    """셀 내용을 text 로 교체. 첫 문단의 첫 run 서식을 물려받는다(서식 보존)."""
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""
    if align is not None:
        p.alignment = align
    # 배분 정렬(distribute)이면 짧은 값이 칸 너비로 늘어난다 — 좌측으로 되돌린다
    if p.alignment == WD_ALIGN_PARAGRAPH.DISTRIBUTE:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def clear_highlight(doc):
    """디렉터가 칠한 형광펜·노란 셀배경 제거 (완성본에 그대로 남는다)"""
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for hl in doc.element.body.iter(ns + "highlight"):
        hl.set(ns + "val", "none")
    for shd in doc.element.body.iter(ns + "shd"):
        if (shd.get(ns + "fill") or "").upper() in ("FFFF00", "FFFF99", "FFF200"):
            shd.set(ns + "fill", "auto")


def build(kind):
    src, out = SRC[kind], OUT[kind]
    shutil.copy(src, out)
    doc = Document(out)

    t_approval, t_info = doc.tables[0], doc.tables[1]

    # ── 1) 결재표: 2행 서명칸에 마커 ─────────────────────────
    hdr = [cell_text(c).strip() for c in t_approval.rows[0].cells]
    sign_row = t_approval.rows[1]
    seen = set()
    for i, label in enumerate(hdr):
        cell = sign_row.cells[i]
        if id(cell._tc) in seen:      # 병합 셀 중복 방지
            continue
        seen.add(id(cell._tc))
        if "작성자" in label:
            set_cell(cell, "《근로자서명》", WD_ALIGN_PARAGRAPH.CENTER)
        elif "원장" in label:
            set_cell(cell, "《원장서명》", WD_ALIGN_PARAGRAPH.CENTER)
        elif "본부" in label:
            set_cell(cell, "《본부서명》", WD_ALIGN_PARAGRAPH.CENTER)

    # ── 2) 신청인 정보표 ────────────────────────────────────
    for row in t_info.rows:
        cells = row.cells
        texts = [cell_text(c).strip() for c in cells]
        for i, t in enumerate(texts):
            nxt = cells[i + 1] if i + 1 < len(cells) else None
            if nxt is None:
                continue
            if t == "성명":
                set_cell(nxt, "{직원명}", WD_ALIGN_PARAGRAPH.CENTER)
            elif t == "생년월일":
                set_cell(nxt, "{생년월일}", WD_ALIGN_PARAGRAPH.CENTER)
            elif t.startswith("소속"):
                set_cell(nxt, "{지점}", WD_ALIGN_PARAGRAPH.CENTER)
            elif t.startswith("직급"):
                set_cell(nxt, "{직급}", WD_ALIGN_PARAGRAPH.CENTER)
            elif t == "신청구분":
                set_cell(nxt, "{체크_12주이내} 임신 12주 이내 해당  /  "
                              "{체크_32주이후} 임신 32주 이후 해당  /  {체크_기타} 기타",
                         WD_ALIGN_PARAGRAPH.LEFT)
            elif t == "출산예정일":
                set_cell(nxt, "{출산예정일}", WD_ALIGN_PARAGRAPH.CENTER)
            elif t.replace(" ", "").startswith("근로시간단축기간"):
                set_cell(nxt, "{단축시작일}  ~  {단축종료일}", WD_ALIGN_PARAGRAPH.CENTER)
            elif t.replace(" ", "").startswith("근무개시"):
                set_cell(nxt, "{근무개시시각}  ~  {근무종료시각}", WD_ALIGN_PARAGRAPH.CENTER)

    # ── 3) 표 밖 문단: 작성일 · 신청인 서명 ──────────────────
    for p in doc.paragraphs:
        txt = p.text
        if re.search(r"20\s+년\s+월\s+일", txt):
            set_para(p, "{작성일}", WD_ALIGN_PARAGRAPH.CENTER)
        elif txt.strip().startswith("신청인"):
            set_para(p, "신청인 : {직원명}  《근로자서명》  (인)", WD_ALIGN_PARAGRAPH.RIGHT)

    clear_highlight(doc)
    doc.save(out)
    return out


def set_para(p, text, align=None):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    if align is not None:
        p.alignment = align


for kind in ("manager", "director"):
    out = build(kind)
    # 결과 검증
    doc = Document(out)
    body = "\n".join(p.text for p in doc.paragraphs)
    for t in doc.tables:
        for row in t.rows:
            for c in row.cells:
                body += "\n" + cell_text(c)
    fields = sorted(set(re.findall(r"\{([^}]+)\}", body)))
    markers = sorted(set(re.findall(r"《([^》]+)》", body)))
    print(f"[{kind}] {out}")
    print("   필드  :", ", ".join(fields))
    print("   마커  :", ", ".join(markers))
