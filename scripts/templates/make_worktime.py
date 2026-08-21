# 근로시간 변경 신청서 — 필드 심기
#
# 구조: 제목표(내부에 결재표 중첩: 원장|본부) + 본문표
#   소속(에듀플렉스 __직영점) 직급 / 성명 사번 / 변경적용일 /
#   주 소정근로시간 변경 전·후 / 변경 사유 / 동의사항 / 작성일+신청인 서명
#
# 심는 것
#   자동    {직원명} {직급} {사원번호} {지점} {작성일}
#   관리자  {변경적용일} {변경전근로시간} {변경후근로시간}
#   직원    {변경사유}  ("사유" 화이트리스트 — 서명 시 직접 입력)
#   서명    《근로자서명》 《원장서명》 《본부서명》
import re, shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"C:/Users/N-88/Downloads/근로시간 변경 신청서 양식.docx"
OUT = "근로시간변경신청서_템플릿.docx"


def cell_text(c):
    return "".join(p.text for p in c.paragraphs)


def set_cell(cell, text, align=None):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""
    if align is not None:
        p.alignment = align
    if p.alignment == WD_ALIGN_PARAGRAPH.DISTRIBUTE:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_para(p, text, align=None):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    if align is not None:
        p.alignment = align


def clear_highlight(doc):
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for hl in doc.element.body.iter(ns + "highlight"):
        hl.set(ns + "val", "none")
    for shd in doc.element.body.iter(ns + "shd"):
        if (shd.get(ns + "fill") or "").upper() in ("FFFF00", "FFFF99", "FFF200"):
            shd.set(ns + "fill", "auto")


shutil.copy(SRC, OUT)
doc = Document(OUT)

# ── 1) 결재표(중첩 표) — 원장/본부 라벨 아래 빈 행에 마커 ──
# doc.tables 는 최상위 표만 주므로, 표 안의 표는 셀에서 찾는다
def iter_all_tables(doc):
    for t in doc.tables:
        yield t
        for row in t.rows:
            for c in row.cells:
                for nt in c.tables:
                    yield nt

approval_done = False
for t in iter_all_tables(doc):
    rows = t.rows
    if len(rows) < 2 or approval_done:
        continue
    hdr = [cell_text(c).strip() for c in rows[0].cells]
    if "원장" in hdr and "본부" in hdr:
        sign_row = rows[1]
        seen = set()
        for i, label in enumerate(hdr):
            cell = sign_row.cells[i]
            if id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            if "원장" in label:
                set_cell(cell, "《원장서명》", WD_ALIGN_PARAGRAPH.CENTER)
            elif "본부" in label:
                set_cell(cell, "《본부서명》", WD_ALIGN_PARAGRAPH.CENTER)
        approval_done = True

# ── 2) 본문 표 ───────────────────────────────────────────────
for t in iter_all_tables(doc):
    for row in t.rows:
        cells = row.cells
        texts = [cell_text(c).strip() for c in cells]
        joined = "".join(texts)
        norm = lambda x: "".join(x.split())  # 줄바꿈·공백 제거 비교 (라벨에 개행 있는 셀 대응)
        for i, txt in enumerate(texts):
            txt = norm(txt)
            nxt = cells[i + 1] if i + 1 < len(cells) else None
            if txt == "소속" and nxt is not None and "에듀플렉스" in cell_text(nxt):
                set_cell(nxt, "넥스큐브코퍼레이션㈜ 에듀플렉스 {지점} 직영점", WD_ALIGN_PARAGRAPH.CENTER)
            elif txt == "직급" and nxt is not None and not cell_text(nxt).strip():
                set_cell(nxt, "{직급}", WD_ALIGN_PARAGRAPH.CENTER)
            elif txt == "성명" and nxt is not None and not cell_text(nxt).strip():
                set_cell(nxt, "{직원명}", WD_ALIGN_PARAGRAPH.CENTER)
            elif txt == "사번" and nxt is not None and not cell_text(nxt).strip():
                set_cell(nxt, "{사원번호}", WD_ALIGN_PARAGRAPH.CENTER)
            elif txt == "변경적용일" and nxt is not None:
                set_cell(nxt, "{변경적용일} ~", WD_ALIGN_PARAGRAPH.CENTER)
            elif txt == "변경사유" and nxt is not None and not cell_text(nxt).strip():
                # 병합 반복 셀(라벨 자신)을 덮지 않게 빈 셀에만 넣는다
                set_cell(nxt, "{변경사유}", WD_ALIGN_PARAGRAPH.LEFT)
        # 주 소정근로시간 행: 라벨 뒤 빈칸 2개 = 변경 전 / 변경 후
        if texts and norm(texts[0]) == "주소정근로시간":
            empties = [c for c in cells[1:] if not cell_text(c).strip()]
            if len(empties) >= 2:
                set_cell(empties[0], "{변경전근로시간}", WD_ALIGN_PARAGRAPH.CENTER)
                set_cell(empties[1], "{변경후근로시간}", WD_ALIGN_PARAGRAPH.CENTER)
        # 작성일 + 신청인 서명 (한 셀에 여러 문단)
        if "신청인" in joined and re.search(r"20\s+년", joined):
            target = next(c for c in cells if "신청인" in cell_text(c))
            for p in target.paragraphs:
                if re.search(r"20\s+년\s+월\s+일", p.text):
                    set_para(p, "{작성일}", WD_ALIGN_PARAGRAPH.CENTER)
                elif "신청인" in p.text:
                    set_para(p, "신청인(근로자) : {직원명}  《근로자서명》  (인)", WD_ALIGN_PARAGRAPH.RIGHT)

clear_highlight(doc)
doc.save(OUT)

# ── 검증 ────────────────────────────────────────────────────
doc = Document(OUT)
body = ""
for t in iter_all_tables(doc):
    for row in t.rows:
        for c in row.cells:
            body += "\n" + cell_text(c)
for p in doc.paragraphs:
    body += "\n" + p.text
fields = sorted(set(re.findall(r"\{([^}]+)\}", body)))
markers = sorted(set(re.findall(r"《([^》]+)》", body)))
print("필드:", ", ".join(fields))
print("마커:", ", ".join(markers))
