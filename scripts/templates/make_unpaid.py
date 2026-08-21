# 무급휴가 사용 동의서 — 필드 심기
#
# 구조(디렉터 수정본): 인적사항 표 + 신청일자 문단 + 동의 내용 ①~⑦ +
#   설명 확인 ☐ 2개(직원이 직접 체크) + 작성일 + 근로자/본부/원장 서명 줄
#
# 심는 것
#   자동    {직원명} {지점} {직급} {연락처} {작성일}
#   관리자  {무급휴가시작일} {무급휴가종료일} {휴가일수}
#   직원    {확인_불이익 설명 확인} {확인_거부권 안내 확인}  ← 서명 시 직접 체크(필수)
#   서명    《근로자서명》 《원장서명》 《본부서명》
import re, shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"C:/Users/N-88/Downloads/무급휴가_사용_동의서.docx"
OUT = "무급휴가_사용_동의서_템플릿.docx"


def cell_text(c):
    return "".join(p.text for p in c.paragraphs)


def set_cell(cell, text, align=None):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""
    if align is not None:
        p.alignment = align
    if p.alignment == WD_ALIGN_PARAGRAPH.DISTRIBUTE:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_para(p, text, align=None):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    if align is not None:
        p.alignment = align


def clear_highlight(doc):
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for hl in doc.element.body.iter(ns + "highlight"):
        hl.set(ns + "val", "none")
    for shd in doc.element.body.iter(ns + "shd"):
        if (shd.get(ns + "fill") or "").upper() in ("FFFF00", "FFFF99", "FFF200"):
            shd.set(ns + "fill", "auto")


shutil.copy(SRC, OUT)
doc = Document(OUT)

# ── 1) 인적사항 표 ──────────────────────────────────────────
norm = lambda x: "".join(x.split())
for t in doc.tables:
    for row in t.rows:
        cells = row.cells
        texts = [cell_text(c) for c in cells]
        for i, txt in enumerate(texts):
            label = norm(txt)
            nxt = cells[i + 1] if i + 1 < len(cells) else None
            if nxt is None or cell_text(nxt).strip():
                continue  # 빈 값 칸에만 넣는다 (병합 반복 셀 라벨 덮기 방지)
            if label in ("소속/지점", "소속"):
                set_cell(nxt, "{지점}", WD_ALIGN_PARAGRAPH.CENTER)
            elif label == "직급":
                set_cell(nxt, "{직급}", WD_ALIGN_PARAGRAPH.CENTER)
            elif label == "성명":
                set_cell(nxt, "{직원명}", WD_ALIGN_PARAGRAPH.CENTER)
            elif label == "연락처":
                set_cell(nxt, "{연락처}", WD_ALIGN_PARAGRAPH.CENTER)

# ── 2) 문단들 ───────────────────────────────────────────────
for p in doc.paragraphs:
    txt = p.text
    if txt.strip().startswith("2. 신청 일자"):
        set_para(p, "2. 신청 일자 :  {무급휴가시작일}  ~  {무급휴가종료일}  (휴일제외 {휴가일수}일간)")
    elif txt.strip().startswith("☐") and "불이익 사항" in txt:
        set_para(p, "{확인_불이익 설명 확인}  불이익 사항(급여·주휴수당·차년도 연차·4대보험)에 대하여 설명을 들었습니다.")
    elif txt.strip().startswith("☐") and "거부" in txt:
        set_para(p, "{확인_거부권 안내 확인}  동의를 거부할 수 있고, 거부하더라도 인사상 불이익이 없다는 점을 안내받았습니다.")
    elif txt.strip().startswith("작성일"):
        set_para(p, "작성일     {작성일}")
    elif "동의자(근로자)" in txt:
        set_para(p, "동의자(근로자)     성명 : {직원명}   《근로자서명》   (서명 또는 인)", WD_ALIGN_PARAGRAPH.RIGHT)
    elif txt.strip().startswith("본부") and "직영사업본부" in txt:
        set_para(p, "본부     직영사업본부   《본부서명》   (서명 또는 인)", WD_ALIGN_PARAGRAPH.RIGHT)
    elif txt.strip().startswith("원장") and "서명 또는 인" in txt:
        set_para(p, "원장     《원장서명》   (서명 또는 인)", WD_ALIGN_PARAGRAPH.RIGHT)

clear_highlight(doc)
doc.save(OUT)

# ── 검증 ────────────────────────────────────────────────────
doc = Document(OUT)
body = "\n".join(p.text for p in doc.paragraphs)
for t in doc.tables:
    for row in t.rows:
        for c in row.cells:
            body += "\n" + cell_text(c)
fields = sorted(set(re.findall(r"\{([^}]+)\}", body)))
markers = sorted(set(re.findall(r"《([^》]+)》", body)))
print("필드:", ", ".join(fields))
print("마커:", ", ".join(markers))
print("남은 ☐:", body.count("☐"))
