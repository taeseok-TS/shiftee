# 유급휴가 선사용 동의서 — 필드 심기
#
# 구조: 제목표(결재표 중첩: 원장|본부) + 본문표
#   소속/직급/성명/사번 + 발생예정일(○ 2줄)·발생예정 총일수 +
#   선사용기간(~, 총일수) + 사유 + 동의사항 + 작성일·신청인 서명
#
# 심는 것
#   자동    {직원명} {직급} {사원번호} {지점} {작성일}
#   관리자  {발생예정일} {추가 발생예정일} {발생예정휴가일수}
#           {선사용시작일} {선사용종료일} {선사용일수}
#   직원    {선사용사유}  (사유 화이트리스트 — 서명 시 입력)
#   서명    《근로자서명》 《원장서명》 《본부서명》
import re, shutil
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH

SRC = r"C:/Users/N-88/Downloads/유급휴가 선사용 동의서.docx"
OUT = "유급휴가_선사용_동의서_템플릿.docx"


def cell_text(c):
    return "".join(p.text for p in c.paragraphs)


def set_cell(cell, text, align=None):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    for extra in cell.paragraphs[1:]:
        for r in extra.runs:
            r.text = ""
    if align is not None:
        p.alignment = align
    if p.alignment == WD_ALIGN_PARAGRAPH.DISTRIBUTE:
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT


def set_para(p, text, align=None):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)
    if align is not None:
        p.alignment = align


def clear_highlight(doc):
    ns = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    for hl in doc.element.body.iter(ns + "highlight"):
        hl.set(ns + "val", "none")
    for shd in doc.element.body.iter(ns + "shd"):
        if (shd.get(ns + "fill") or "").upper() in ("FFFF00", "FFFF99", "FFF200"):
            shd.set(ns + "fill", "auto")


def iter_all_tables(doc):
    for t in doc.tables:
        yield t
        for row in t.rows:
            for c in row.cells:
                for nt in c.tables:
                    yield nt


shutil.copy(SRC, OUT)
doc = Document(OUT)
norm = lambda x: "".join(x.split())

# ── 1) 결재표(중첩) ─────────────────────────────────────────
approval_done = False
for t in iter_all_tables(doc):
    rows = t.rows
    if len(rows) < 2 or approval_done:
        continue
    hdr = [cell_text(c).strip() for c in rows[0].cells]
    if "원장" in hdr and "본부" in hdr:
        sign_row = rows[1]
        seen = set()
        for i, label in enumerate(hdr):
            cell = sign_row.cells[i]
            if id(cell._tc) in seen:
                continue
            seen.add(id(cell._tc))
            if "원장" in label:
                set_cell(cell, "《원장서명》", WD_ALIGN_PARAGRAPH.CENTER)
            elif "본부" in label:
                set_cell(cell, "《본부서명》", WD_ALIGN_PARAGRAPH.CENTER)
        approval_done = True

# ── 2) 본문 표 ──────────────────────────────────────────────
for t in iter_all_tables(doc):
    for row in t.rows:
        cells = row.cells
        texts = [cell_text(c).strip() for c in cells]
        joined = "".join(texts)
        for i, raw in enumerate(texts):
            txt = norm(raw)
            nxt = cells[i + 1] if i + 1 < len(cells) else None
            if txt == "소속" and nxt is not None and "에듀플렉스" in cell_text(nxt):
                set_cell(nxt, "넥스큐브코퍼레이션㈜ 에듀플렉스 {지점} 직영점", WD_ALIGN_PARAGRAPH.CENTER)
            elif txt == "직급" and nxt is not None and not cell_text(nxt).strip():
                set_cell(nxt, "{직급}", WD_ALIGN_PARAGRAPH.CENTER)
            elif txt == "성명" and nxt is not None and not cell_text(nxt).strip():
                set_cell(nxt, "{직원명}", WD_ALIGN_PARAGRAPH.CENTER)
            elif txt == "사번" and nxt is not None and not cell_text(nxt).strip():
                set_cell(nxt, "{사원번호}", WD_ALIGN_PARAGRAPH.CENTER)
            elif txt == "유급휴가발생예정일" and nxt is not None and "년" in cell_text(nxt):
                # ○ 두 줄 — 문단별로 교체 (둘째는 없으면 빈칸으로 나온다)
                ps = [p for p in nxt.paragraphs if p.text.strip()]
                if ps:
                    set_para(ps[0], "○  {발생예정일}", WD_ALIGN_PARAGRAPH.LEFT)
                if len(ps) > 1:
                    set_para(ps[1], "○  {추가 발생예정일}", WD_ALIGN_PARAGRAPH.LEFT)
            elif txt == "발생예정유급휴가" and nxt is not None:
                set_cell(nxt, "총  {발생예정휴가일수} 일", WD_ALIGN_PARAGRAPH.CENTER)
            elif txt == "유급휴가선사용기간" and nxt is not None and "년" in cell_text(nxt):
                set_cell(nxt, "{선사용시작일}  ~  {선사용종료일}   (총 {선사용일수}일)", WD_ALIGN_PARAGRAPH.CENTER)
            elif txt == "사유" and nxt is not None and not cell_text(nxt).strip():
                set_cell(nxt, "{선사용사유}", WD_ALIGN_PARAGRAPH.LEFT)
        # 작성일 + 신청인 서명 (한 셀 여러 문단)
        if "신청인" in joined and re.search(r"20\s+년", joined):
            target = next(c for c in cells if "신청인" in cell_text(c))
            for p in target.paragraphs:
                if re.search(r"20\s+년\s+월\s+일", p.text):
                    set_para(p, "{작성일}", WD_ALIGN_PARAGRAPH.CENTER)
                elif "신청인" in p.text:
                    set_para(p, "신청인(근로자) : {직원명}  《근로자서명》  (서명 / 인)", WD_ALIGN_PARAGRAPH.RIGHT)

clear_highlight(doc)
doc.save(OUT)

# ── 검증 ────────────────────────────────────────────────────
doc = Document(OUT)
body = ""
for t in iter_all_tables(doc):
    for row in t.rows:
        for c in row.cells:
            body += "\n" + cell_text(c)
for p in doc.paragraphs:
    body += "\n" + p.text
fields = sorted(set(re.findall(r"\{([^}]+)\}", body)))
markers = sorted(set(re.findall(r"《([^》]+)》", body)))
print("필드:", ", ".join(fields))
print("마커:", ", ".join(markers))
